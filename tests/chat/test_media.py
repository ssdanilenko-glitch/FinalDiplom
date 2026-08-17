"""Тесты конвертации медиа в content-part."""

import base64
from io import BytesIO
from unittest.mock import AsyncMock, MagicMock

import pytest
from docx import Document as DocxDocument
from pypdf import PdfWriter

from app.chat.media import extract_docx_text, extract_pdf_text, media_to_part


class FakeUploadFile:
    """Минимальный двойник Starlette UploadFile."""

    def __init__(
        self, content_type: str, data: bytes, filename: str = "file"
    ):
        self.content_type = content_type
        self.filename = filename
        self.size = len(data)
        self._data = data

    async def read(self) -> bytes:
        return self._data


@pytest.mark.asyncio
async def test_image_to_part_returns_base64_data_uri():
    data = b"\x89PNG\r\n\x1a\nfake-bytes"
    f = FakeUploadFile("image/png", data, "test.png")
    llm = MagicMock()  # не должен дёргаться для image
    part = await media_to_part(f, llm)
    assert part["type"] == "image_url"
    assert part["image_url"]["url"].startswith("data:image/png;base64,")
    b64 = part["image_url"]["url"].split(",", 1)[1]
    assert base64.b64decode(b64) == data


@pytest.mark.asyncio
async def test_audio_to_part_calls_whisper():
    data = b"OggS\x00\x02fake-ogg-bytes"
    f = FakeUploadFile("audio/ogg", data, "voice.ogg")

    llm = MagicMock()
    llm.audio.transcriptions.create = AsyncMock(
        return_value=MagicMock(text="Привет, как дела"),
    )

    part = await media_to_part(f, llm)
    assert part["type"] == "text"
    assert "пользователь сказал голосом" in part["text"]
    assert "Привет, как дела" in part["text"]
    llm.audio.transcriptions.create.assert_awaited_once()


@pytest.mark.asyncio
async def test_application_ogg_treated_as_audio():
    """Telegram voice иногда летит как application/ogg, не audio/*."""
    data = b"OggS\x00\x02"
    f = FakeUploadFile("application/ogg", data, "v.ogg")
    llm = MagicMock()
    llm.audio.transcriptions.create = AsyncMock(
        return_value=MagicMock(text="Тест"),
    )
    part = await media_to_part(f, llm)
    assert part["type"] == "text"
    assert "Тест" in part["text"]


@pytest.mark.asyncio
async def test_pdf_to_part(tmp_path):
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    pdf_path = tmp_path / "test.pdf"
    with open(pdf_path, "wb") as fp:
        writer.write(fp)

    data = pdf_path.read_bytes()
    f = FakeUploadFile("application/pdf", data, "test.pdf")
    llm = MagicMock()
    part = await media_to_part(f, llm)
    assert part["type"] == "text"
    assert "документ PDF" in part["text"]


@pytest.mark.asyncio
async def test_docx_to_part(tmp_path):
    doc = DocxDocument()
    doc.add_paragraph("Hello world")
    doc.add_paragraph("Второй параграф")
    out = tmp_path / "test.docx"
    doc.save(out)

    data = out.read_bytes()
    mime = (
        "application/vnd.openxmlformats-officedocument."
        "wordprocessingml.document"
    )
    f = FakeUploadFile(mime, data, "test.docx")
    llm = MagicMock()
    part = await media_to_part(f, llm)
    assert part["type"] == "text"
    assert "документ DOCX" in part["text"]
    assert "Hello world" in part["text"]
    assert "Второй параграф" in part["text"]


@pytest.mark.asyncio
async def test_unknown_mime_raises():
    f = FakeUploadFile("application/octet-stream", b"\x00\x01\x02")
    llm = MagicMock()
    with pytest.raises(ValueError, match="Unsupported"):
        await media_to_part(f, llm)


def test_extract_pdf_text_handles_scan_placeholder():
    """Создадим PDF с 5 пустыми страницами — функция должна вернуть placeholder."""
    writer = PdfWriter()
    for _ in range(5):
        writer.add_blank_page(width=72, height=72)
    buf = BytesIO()
    writer.write(buf)
    text = extract_pdf_text(buf.getvalue())
    assert "скан" in text or "OCR" in text


def test_extract_docx_text_with_table(tmp_path):
    doc = DocxDocument()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "C"
    table.cell(1, 1).text = "D"
    out = tmp_path / "t.docx"
    doc.save(out)

    text = extract_docx_text(out.read_bytes())
    assert "A | B" in text
    assert "C | D" in text
